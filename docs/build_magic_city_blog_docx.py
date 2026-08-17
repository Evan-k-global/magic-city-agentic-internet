from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "magic-city-application-layer-expanded.md"
OUTPUT = ROOT / "Magic-City-The-Internet-Application-Layer.docx"
DIAGRAM_DIR = ROOT / "diagrams"

NAVY = RGBColor(14, 27, 46)
BLUE = RGBColor(31, 104, 151)
CYAN = RGBColor(32, 186, 208)
PINK = RGBColor(222, 68, 159)
INK = RGBColor(28, 35, 45)
MUTED = RGBColor(92, 104, 119)
LIGHT = "EEF3F8"
RULE = "D8E2EC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=RULE, size=6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def restart_numbered_lists(doc: Document) -> None:
    numbering = doc.part.numbering_part.element
    style = doc.styles["List Number"]
    style_num_id = style.element.pPr.numPr.numId.val
    source_num = next(
        node for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == int(style_num_id)
    )
    abstract_id = source_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    next_num_id = max(existing_ids) + 1
    active_num_id = None
    previous_was_numbered = False

    for paragraph in doc.paragraphs:
        is_numbered = paragraph.style.name == "List Number"
        if not is_numbered:
            previous_was_numbered = False
            active_num_id = None
            continue
        if not previous_was_numbered:
            active_num_id = next_num_id
            next_num_id += 1
            num = OxmlElement("w:num")
            num.set(qn("w:numId"), str(active_num_id))
            abstract = OxmlElement("w:abstractNumId")
            abstract.set(qn("w:val"), str(abstract_id))
            num.append(abstract)
            override = OxmlElement("w:lvlOverride")
            override.set(qn("w:ilvl"), "0")
            start = OxmlElement("w:startOverride")
            start.set(qn("w:val"), "1")
            override.append(start)
            num.append(override)
            numbering.append(num)

        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        ilvl = num_pr.find(qn("w:ilvl"))
        if ilvl is None:
            ilvl = OxmlElement("w:ilvl")
            num_pr.append(ilvl)
        ilvl.set(qn("w:val"), "0")
        num_id = num_pr.find(qn("w:numId"))
        if num_id is None:
            num_id = OxmlElement("w:numId")
            num_pr.append(num_id)
        num_id.set(qn("w:val"), str(active_num_id))
        previous_was_numbered = True


def apply_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_rich_text(paragraph, text: str, base_size=11, base_color=INK) -> None:
    token_pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    cursor = 0
    for match in token_pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            apply_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            apply_font(run, size=base_size, color=base_color, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            apply_font(run, size=base_size, color=base_color, italic=True)
        else:
            run = paragraph.add_run(token[1:-1])
            apply_font(run, name="Courier New", size=base_size - 0.5, color=NAVY)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        apply_font(run, size=base_size, color=base_color)


def add_caption(doc, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = False
    run = p.add_run(caption)
    apply_font(run, size=9, color=MUTED, italic=True)


def add_diagram(doc, alt: str, relative_path: str, figure_no: int) -> None:
    image_path = ROOT / relative_path
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    picture = run.add_picture(str(image_path), width=Inches(6.35))
    picture._inline.docPr.set("descr", alt)
    picture._inline.docPr.set("title", f"Figure {figure_no}")
    add_caption(doc, f"Figure {figure_no}. {alt}")


def add_code_block(doc, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    prevent_row_split(table.rows[0])
    set_cell_shading(cell, "F4F7FA")
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    set_table_borders(table, color="D3DEE8", size=5)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for i, line in enumerate(lines):
        if i:
            p.add_run().add_break()
        run = p.add_run(line)
        apply_font(run, name="Courier New", size=8.7, color=NAVY)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.22
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for style_name, size, color, before, after in (
        ("Heading 1", 17, NAVY, 18, 8),
        ("Heading 2", 13.5, BLUE, 14, 6),
        ("Heading 3", 11.5, NAVY, 10, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    quote = styles["Quote"]
    quote.font.name = "Calibri"
    quote.font.size = Pt(11)
    quote.font.italic = True
    quote.font.color.rgb = NAVY
    quote.paragraph_format.left_indent = Inches(0.3)
    quote.paragraph_format.right_indent = Inches(0.3)
    quote.paragraph_format.space_before = Pt(5)
    quote.paragraph_format.space_after = Pt(10)

    if "Figure Caption" not in [s.name for s in styles]:
        styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("MAGIC CITY | AGENTIC INTERNET")
    apply_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Magic City")
    apply_font(run, size=8, color=MUTED)
    run = p.add_run("   |   ")
    apply_font(run, size=8, color=CYAN)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def add_cover(doc: Document, title: str, abstract: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("MAGIC CITY")
    apply_font(run, size=11, color=PINK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(title)
    apply_font(run, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(abstract)
    apply_font(run, size=13.5, color=BLUE, bold=True)

    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Inches(2.15), Inches(2.15), Inches(2.2)]
    labels = [
        ("IDENTITY", "Mission-bound holder"),
        ("PAYMENT", "Credits + Base USDC"),
        ("STATE", "Receipts + Zeko"),
    ]
    for idx, (cell, width, (label, value)) in enumerate(zip(table.rows[0].cells, widths, labels)):
        cell.width = width
        set_cell_shading(cell, "F3F7FA")
        set_cell_margins(cell, top=130, start=130, bottom=130, end=130)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label)
        apply_font(r, size=8, color=PINK if idx == 0 else CYAN, bold=True)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        apply_font(r, size=9.5, color=NAVY, bold=True)
    set_table_borders(table, color="DCE6EF", size=5)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("A technical essay on bounded agent authority, open execution markets, and verifiable internet work.")
    apply_font(r, size=10.5, color=MUTED, italic=True)


def parse_markdown(doc: Document, lines: list[str]) -> None:
    title = lines[0].lstrip("# ").strip()
    abstract_line = next(line for line in lines if line.startswith("**Abstract.**"))
    abstract = abstract_line.replace("**Abstract.**", "").strip()
    add_cover(doc, title, abstract)

    i = 1
    figure_no = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        line = raw.strip()
        if not line or line.startswith("**Abstract.**"):
            i += 1
            continue
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue
        image = re.match(r"!\[(.+?)\]\((.+?)\)", line)
        if image:
            figure_no += 1
            add_diagram(doc, image.group(1), image.group(2), figure_no)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
            i += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
            i += 1
            continue
        if re.match(r"^\d+\. ", line):
            content = re.sub(r"^\d+\. ", "", line)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.18
            add_rich_text(p, content)
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.18
            add_rich_text(p, line[2:])
            i += 1
            continue
        if line.startswith("> "):
            p = doc.add_paragraph(style="Quote")
            add_rich_text(p, line[2:])
            i += 1
            continue

        paragraph_parts = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith(("#", "- ", "> ", "```", "![")) or re.match(r"^\d+\. ", nxt):
                break
            paragraph_parts.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_rich_text(p, " ".join(paragraph_parts))


def build() -> None:
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    parse_markdown(doc, lines)
    restart_numbered_lists(doc)

    core = doc.core_properties
    core.title = "The Internet Never Had an Application Layer. We Just Built One."
    core.subject = "Magic City, Agent Mission-Bound Auth, Zeko, and SantaClawz"
    core.author = "Magic City"
    core.keywords = "Magic City, AMBA, Zeko, SantaClawz, x402, agents"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
